import docx
import re
from typing import Dict, List, Tuple

class DocProcessor:
    def __init__(self, file_path: str):
        """初始化文档处理器

        Args:
            file_path (str): Word文档路径
        """
        self.file_path = file_path
        self.doc = docx.Document(file_path)

    def extract_amount_sentences(self) -> List[Tuple[str, float]]:
        """提取包含'万元'的句子及其金额

        Returns:
            List[Tuple[str, float]]: 包含'万元'的句子及其金额（单位：元）列表
        """
        sentences_with_amount = []
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if '万元' in text:
                matches = re.finditer(r'[^。！？.，"”]*?(\d+)万元[^。，！？.]*', text)
                for match in matches:
                    sentence = match.group().strip().rstrip('。！？.，')
                    amount = float(match.group(1)) * 10000  # 转换为元
                    sentences_with_amount.append((sentence, amount))
        return sentences_with_amount

    def parse_expense_table(self) -> Tuple[Dict[str, float], float, float]:
        """解析费用明细表格

        Returns:
            Tuple[Dict[str, float], float, float]: (费用明细字典, 计算总金额, 表格汇总金额)
        """
        expense_details = {}
        total_amount = 0.0
        table_total = 0.0

        for table in self.doc.tables:
            if self._is_expense_table(table):
                data_rows = table.rows[1:-1]
                if len(table.rows) > 2:
                    last_row = table.rows[-1]
                    last_row_amount = last_row.cells[-1].text.strip()
                    total_match = re.search(r'\d+(\.\d+)?', last_row_amount)
                    if total_match:
                        table_total = float(total_match.group())

                processed_cells = set()
                for row in data_rows:
                    cell_key = (row.cells[0]._tc, row.cells[-1]._tc)
                    if cell_key in processed_cells:
                        continue

                    content = row.cells[0].text.strip()
                    amount_text = row.cells[-1].text.strip()
                    amount_match = re.search(r'\d+(\.\d+)?', amount_text)
                    if amount_match:
                        amount = float(amount_match.group())
                        expense_details[content] = amount
                        total_amount += amount
                        processed_cells.add(cell_key)

        return expense_details, total_amount, table_total

    def _is_expense_table(self, table) -> bool:
        if not table.rows:
            return False
            
        header_text = ' '.join(cell.text.strip() for cell in table.rows[0].cells)
        if '费用明细' in header_text:
            return True

        current_element = table._element
        previous_element = current_element.getprevious()
        
        while previous_element is not None:
            if previous_element.tag.endswith('p'):
                paragraph_text = previous_element.xpath('.//w:t')
                if paragraph_text:
                    text = ''.join([t.text for t in paragraph_text if t.text])
                    if '费用明细' in text:
                        return True
                break
            previous_element = previous_element.getprevious()

        return False

def process_doc(file_path: str):
    try:
        processor = DocProcessor(file_path)
        
        # 1. 提取包含'万元'的句子及其金额
        print('一、文档中包含的金额信息：')
        sentences_with_amount = processor.extract_amount_sentences()
        for sentence, amount in sentences_with_amount:
            print(f'- {sentence}')

        # 2. 解析费用明细表格
        print('\n二、费用明细表格解析结果：')
        expense_details, total_amount, table_total = processor.parse_expense_table()
        if expense_details:
            for content, amount in expense_details.items():
                print(f'- {content}: {amount:,.2f}元')
            print(f'\n计算总金额：{total_amount:,.2f}元')
            if table_total > 0:
                print(f'表格汇总金额：{table_total:,.2f}元')
                diff = abs(total_amount - table_total)
                if diff < 0.01:
                    print('计算总金额与表格汇总金额一致')
                else:
                    print(f'计算总金额与表格汇总金额相差：{diff:,.2f}元')

        # 3. 金额对比分析
        print('\n三、金额对比分析：')
        if sentences_with_amount and expense_details:
            print('文档中提到的金额与费用明细表对比：')
            for sentence, amount in sentences_with_amount:
                diff = abs(amount - total_amount)
                if diff < 0.01:
                    print(f'- {sentence}中的金额（{amount:,.2f}元）与费用明细表总金额一致')
                else:
                    print(f'- {sentence}中的金额（{amount:,.2f}元）与费用明细表总金额（{total_amount:,.2f}元）相差{diff:,.2f}元')

    except Exception as e:
        print(f'处理文档时出错：{str(e)}')

if __name__ == '__main__':
    doc_path = input('请输入Word文档路径：')
    process_doc(doc_path)